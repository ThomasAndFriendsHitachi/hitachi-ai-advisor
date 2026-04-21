from pathlib import Path
from typing import Any, Dict, List

from .base_extractor import BaseExtractor, ExtractedData
from ..models import IngestionOptions
from ..parsers import TableParser
from ..utils import TextUtils


class DocxExtractor(BaseExtractor):
    def extract(self, path: Path, options: IngestionOptions) -> ExtractedData:
        del options
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required for .docx files. Install with: pip install python-docx"
            ) from exc

        doc = Document(str(path))
        all_text_lines: List[str] = []
        sections: Dict[str, List[str]] = {}
        current_heading = "__body__"
        sections[current_heading] = []

        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading" in style_name:
                current_heading = text
                sections.setdefault(current_heading, [])
            else:
                sections.setdefault(current_heading, []).append(text)

            all_text_lines.append(text)

        raw_tables: List[List[List[str]]] = []
        for table in doc.tables:
            table_rows: List[List[str]] = []
            for row in table.rows:
                table_rows.append([TextUtils.normalize_whitespace(cell.text or "") for cell in row.cells])
            raw_tables.append(table_rows)

        tables = TableParser.normalize_document_tables(raw_tables)

        cleaned_sections: Dict[str, str] = {}
        for heading, lines in sections.items():
            block = TextUtils.normalize_whitespace("\n".join(lines))
            if block:
                cleaned_sections[heading] = block

        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "section_count": len(cleaned_sections),
            "table_count": len(tables),
            "table_row_count": sum(len(table["entries"]) for table in tables),
        }
        full_text = TextUtils.normalize_whitespace("\n".join(all_text_lines))
        return "text", full_text, cleaned_sections or None, tables or None, metadata
